"""
Chunk retriever — primary retrieval engine.
Refactored for ADTC 2026: Safe scoring, crop filtering, and CPU-only offline execution.
"""

import json
import os
import sys
import re
from dataclasses import dataclass
from pathlib import Path
from typing import List, Optional, Dict, Any

import numpy as np

BASE_DIR = Path(__file__).parent
DATASET_DIR = BASE_DIR / "dataset"
INDEX_DIR = BASE_DIR / "index"
MODEL_DIR = BASE_DIR / "model"
EMBED_MODEL = MODEL_DIR / "all-minilm-l6-v2-q8_0.gguf"

TOP_K = 5
MAX_CHUNK_CHARS = 800
RETRIEVAL_MIN_SIMILARITY = 0.35

@dataclass
class Chunk:
    crop: str
    section: str
    text: str
    source: str

_embedder = None

def _get_embedder():
    global _embedder
    if _embedder is None:
        if not EMBED_MODEL.exists():
            raise RuntimeError(f"Embedding model missing at {EMBED_MODEL}")
        try:
            from llama_cpp import Llama
        except ImportError:
            raise ImportError("llama-cpp-python is required.")
        _embedder = Llama(
            model_path=str(EMBED_MODEL),
            embedding=True,
            n_gpu_layers=0,
            n_ctx=512,
            n_threads=4,
            verbose=False,
        )
    return _embedder

def _embed(texts: List[str]) -> np.ndarray:
    e = _get_embedder()
    out = e.create_embedding(texts)
    data = out["data"]
    vectors = np.array([d["embedding"] for d in sorted(data, key=lambda d: d["index"])], dtype="float32")
    norms = np.linalg.norm(vectors, axis=1, keepdims=True)
    return vectors / np.maximum(norms, 1e-9)

def load_crop_docx(path: str) -> List[Chunk]:
    from docx import Document
    doc = Document(path)
    chunks: List[Chunk] = []
    crop_name = Path(path).stem
    current_section: Optional[str] = None
    buffer: List[str] = []

    def flush():
        if current_section and buffer:
            text = " ".join(buffer).strip()
            if text:
                chunks.append(Chunk(
                    crop=crop_name,
                    section=current_section,
                    text=text,
                    source=f"{Path(path).name} — compiled agricultural reference",
                ))

    for para in doc.paragraphs:
        style = (para.style.name if para.style else "").lower()
        text = para.text.strip()
        if not text: continue
        
        if style == "heading 1":
            flush()
            crop_name = text
            current_section = None
            buffer = []
        elif style == "heading 2":
            flush()
            current_section = text
            buffer = []
        elif style == "heading 3":
            buffer.append(f"\n### {text} ###\n")
        else:
            buffer.append(text)
    flush()

    for table in doc.tables:
        header = [cell.text.strip() for cell in table.rows[0].cells]
        is_tool_table = any("Tool" in h for h in header)
        if is_tool_table:
            for row in table.rows[1:]:
                row_text = [f"{header[i]}: {row.cells[i].text.strip()}" for i in range(len(header)) if i < len(row.cells)]
                text = "\n".join(row_text)
                tool_name = row.cells[0].text.strip().split("(")[0].strip()
                chunks.append(Chunk(crop=crop_name, section=tool_name, text=text, source=f"{Path(path).name} — compiled agricultural reference"))
        else:
            table_text = [" | ".join([c.text.strip() for c in r.cells if c.text.strip()]) for r in table.rows]
            table_text = [t for t in table_text if t]
            if table_text:
                chunks.append(Chunk(crop=crop_name, section=current_section or "Data Table", text="\n".join(table_text), source=f"{Path(path).name} — compiled agricultural reference"))
    return chunks

def build_index():
    import glob
    os.makedirs(INDEX_DIR, exist_ok=True)
    all_chunks = []
    for path in sorted(glob.glob(os.path.join(DATASET_DIR, "*.docx"))):
        all_chunks.extend(load_crop_docx(path))
    if not all_chunks:
        raise RuntimeError("No chunks found.")
    texts = [f"{c.crop} {c.section} {c.text}" for c in all_chunks]
    vectors = _embed(texts)
    np.save(os.path.join(INDEX_DIR, "chunk_vectors.npy"), vectors)
    with open(os.path.join(INDEX_DIR, "chunks.json"), "w", encoding="utf-8") as f:
        json.dump([{"crop": c.crop, "section": c.section, "text": c.text, "source": c.source} for c in all_chunks], f, ensure_ascii=False)
    print(f"Indexed {len(all_chunks)} chunks.")

_cached_index = None

def load_index():
    global _cached_index
    if _cached_index is not None: return _cached_index
    vec_path, chunk_path = INDEX_DIR / "chunk_vectors.npy", INDEX_DIR / "chunks.json"
    if not (vec_path.exists() and chunk_path.exists()): return build_index()
    vectors = np.load(vec_path, allow_pickle=False)
    with open(chunk_path, encoding="utf-8") as f:
        chunks = [Chunk(**c) for c in json.load(f)]
    _cached_index = (vectors, chunks)
    return _cached_index

def retrieve(query: str, top_k: int = TOP_K, crop_filter: Optional[str] = None) -> List[Dict[str, Any]]:
    """
    Retrieves chunks with optional hard crop filtering and honest cosine similarity.
    """
    vectors, chunks = load_index()
    q_vec = _embed([query])[0]
    sims = vectors @ q_vec
    
    # Filter by crop if requested
    if crop_filter:
        filter_mask = np.array([c.crop.lower() == crop_filter.lower() or c.crop.lower() == "farm tools" for c in chunks])
        sims[~filter_mask] = -1.0
        
    ranked = np.argsort(-sims)[:top_k]
    results = []
    for i in ranked:
        score = float(sims[i])
        if score < RETRIEVAL_MIN_SIMILARITY: continue
        c = chunks[i]
        results.append({
            "crop": c.crop,
            "section": c.section,
            "text": c.text[:MAX_CHUNK_CHARS],
            "source": c.source,
            "score": round(score, 3)
        })
    return results

if __name__ == "__main__":
    if len(sys.argv) > 1 and sys.argv[1] == "build":
        build_index()
    else:
        q = " ".join(sys.argv[1:]) if len(sys.argv) > 1 else "how to store beans"
        for r in retrieve(q):
            print(f"[{r['crop']} - {r['section']}] (Score: {r['score']})\n{r['text'][:150]}...\n")
