"""
Chunk retriever — primary retrieval engine.

Knowledge base: crop `.docx` guides in dataset/ (Heading 1 = crop name,
Heading 2 = sections: About, History, Planting, Harvest, Crop Rotation,
Pests, Pest Control, Machines and Tools, Care).

Single runtime, single dependency — llama.cpp for BOTH the chat model and
the embeddings (ADTC 2026 offline rules):

  - model/all-minilm-l6-v2-q8_0.gguf  — quantised all-MiniLM-L6-v2 (25 MB)
    run through llama-cpp-python in embedding mode, the same runtime as
    the Qwen chat model in rag_llm.py
  - model/qwen2.5-1.5b-instruct-q4_k_m.gguf — the chat model

  BUILD (dev-time, one-time + whenever dataset/ changes):
    `python chunk_retriever.py build`
    Ingests the docx guides, embeds every section chunk with the GGUF
    MiniLM model, and writes the vectors to index/ (chunk_vectors.npy +
    chunks.json). Needs: numpy, python-docx, llama-cpp-python
    (download_model.sh pulls the .gguf files first). Never runs at chat time.

  RUNTIME (chat time):
    `from chunk_retriever import retrieve` — reads the pre-built index
    and embeds only the incoming question with the same GGUF MiniLM model
    through llama-cpp-python's embedding mode. Retrieve is a dot product
    over the cached chunk vectors. Sub-second, no cloud calls, no network,
    no torch, no sentence-transformers at runtime.

Chunks below RETRIEVAL_MIN_SIMILARITY are dropped — that's the honest
"not in the dataset" signal.
"""

import json
import os
import sys
from dataclasses import dataclass
from pathlib import Path
from typing import List, Optional

import numpy as np

BASE_DIR = Path(__file__).parent
DATASET_DIR = BASE_DIR / "dataset"       # folder of per-crop .docx files
INDEX_DIR = BASE_DIR / "index"           # pre-built vector store
MODEL_DIR = BASE_DIR / "model"           # GGUF embedding + chat models
EMBED_MODEL = MODEL_DIR / "all-minilm-l6-v2-q8_0.gguf"

TOP_K = 3                              # chunks retrieved per question
MAX_CHUNK_CHARS = 500                    # cap so small LLM context windows aren't blown
RETRIEVAL_MIN_SIMILARITY = 0.35          # below this, treat as "not in the dataset"


@dataclass
class Chunk:
    crop: str
    section: str
    text: str
    source: str


# ---------------------------------------------------------------------------
# Embedding backend — llama.cpp in embedding mode (all llama.cpp, all offline)
# ---------------------------------------------------------------------------

_embedder = None


def _get_embedder():
    """Lazily load the GGUF MiniLM model through llama-cpp-python with
    embedding=True so it only computes sentence embeddings (no chat graph)."""
    global _embedder
    if _embedder is None:
        if not EMBED_MODEL.exists():
            raise RuntimeError(
                f"Embedding model missing at {EMBED_MODEL} — "
                "run download_model.sh (or add all-minilm-l6-v2-q8_0.gguf "
                "to model/) then `python chunk_retriever.py build` once."
            )
        try:
            from llama_cpp import Llama
        except ImportError as e:
            raise ImportError(
                "llama-cpp-python is required at runtime (same runtime as the "
                "chat model). Setup: python -m pip install -r requirements.txt"
            ) from e
        _embedder = Llama(
            model_path=str(EMBED_MODEL),
            embedding=True,
            n_gpu_layers=0,   # keep everything on CPU / within the 8 GB budget
            n_ctx=512,
            n_threads=4, # 4 threads is faster for embedding even on 2-core CPUs
            verbose=False,
        )
    return _embedder


def _embed(texts: List[str]) -> np.ndarray:
    """Embed a batch of strings with the GGUF MiniLM model, L2-normalised
    so a dot product equals cosine similarity."""
    e = _get_embedder()
    out = e.create_embedding(texts)
    data = out["data"]
    vectors = np.array([d["embedding"] for d in sorted(data, key=lambda d: d["index"])], dtype="float32")
    norms = np.linalg.norm(vectors, axis=1, keepdims=True)
    return vectors / np.maximum(norms, 1e-9)


# ---------------------------------------------------------------------------
# Build phase (dev-time) — needs llama-cpp-python + the GGUF model in model/
# ---------------------------------------------------------------------------

def load_crop_docx(path: str) -> List[Chunk]:
    """Heading 2 sections under the Heading 1 crop name -> one chunk each."""
    from docx import Document as DocxDocument  # lazy: only at build time

    doc = DocxDocument(path)
    chunks: List[Chunk] = []
    crop_name = Path(path).stem
    current_section: Optional[str] = None
    buffer: List[str] = []

    def flush():
        if current_section and buffer:
            text = " ".join(buffer).strip()
            if text:
                chunks.append(Chunk(
                    crop=crop_name,
                    section=current_section,
                    text=text,
                    source=f"{Path(path).name} — compiled agricultural reference",
                ))

    # Ingest paragraphs
    for para in doc.paragraphs:
        style = (para.style.name if para.style else "").lower()
        text = para.text.strip()
        if not text:
            continue
        
        if style == "heading 1":
            flush()
            crop_name = text
            current_section = None
            buffer = []
        elif style == "heading 2":
            flush()
            current_section = text
            buffer = []
        elif style == "heading 3":
            buffer.append(f"\n### {text} ###\n")
        else:
            buffer.append(text)
    flush()

    # Ingest tables (especially for Farm Tools)
    for table in doc.tables:
        # Check if this is a tool table (header contains "Tool")
        header = [cell.text.strip() for cell in table.rows[0].cells]
        is_tool_table = any("Tool" in h for h in header)
        
        if is_tool_table:
            # For tool tables, each row (except header) is a chunk
            for row in table.rows[1:]:
                row_text = [f"{header[i]}: {row.cells[i].text.strip()}" for i in range(len(header)) if i < len(row.cells)]
                text = "\n".join(row_text)
                # Extract the tool name for the section field
                tool_name = row.cells[0].text.strip().split("(")[0].strip()
                chunks.append(Chunk(
                    crop=crop_name,
                    section=tool_name,
                    text=text,
                    source=f"{Path(path).name} — compiled agricultural reference",
                ))
        else:
            # For other tables, keep as one chunk
            table_text = []
            for row in table.rows:
                row_text = [cell.text.strip() for cell in row.cells if cell.text.strip()]
                if row_text:
                    table_text.append(" | ".join(row_text))
            if table_text:
                chunks.append(Chunk(
                    crop=crop_name,
                    section=current_section or "Data Table",
                    text="\n".join(table_text),
                    source=f"{Path(path).name} — compiled agricultural reference",
                ))

    return chunks


def build_index():
    """Dev-time: ingest dataset/*.docx and embed all chunks to index/."""
    import glob

    os.makedirs(INDEX_DIR, exist_ok=True)
    all_chunks: List[Chunk] = []
    for path in sorted(glob.glob(os.path.join(DATASET_DIR, "*.docx"))):
        all_chunks.extend(load_crop_docx(path))

    if not all_chunks:
        raise RuntimeError(
            f"No chunks found — check {DATASET_DIR}/ has your crop .docx files "
            "(Rice.docx, Maize.docx, ...). Add more guides and re-run build."
        )

    texts = [f"{c.crop} — {c.section}: {c.text}" for c in all_chunks]
    vectors = _embed(texts)

    np.save(os.path.join(INDEX_DIR, "chunk_vectors.npy"), vectors)
    with open(os.path.join(INDEX_DIR, "chunks.json"), "w", encoding="utf-8") as f:
        json.dump(
            [{"crop": c.crop, "section": c.section, "text": c.text, "source": c.source}
             for c in all_chunks],
            f, ensure_ascii=False,
        )
    print(f"Indexed {len(all_chunks)} chunks from {DATASET_DIR}/")
    return vectors, all_chunks


# ---------------------------------------------------------------------------
# Runtime phase — pre-built index + llama.cpp embedding (no torch)
# ---------------------------------------------------------------------------

_cached_index = None

def load_index():
    """Load index from disk, caching in memory to avoid redundant IO."""
    global _cached_index
    if _cached_index is not None:
        return _cached_index
        
    vec_path = os.path.join(INDEX_DIR, "chunk_vectors.npy")
    chunk_path = os.path.join(INDEX_DIR, "chunks.json")
    if not (os.path.exists(vec_path) and os.path.exists(chunk_path)):
        return build_index()
        
    vectors = np.load(vec_path, allow_pickle=False)
    with open(chunk_path, encoding="utf-8") as f:
        chunks = [Chunk(**c) for c in json.load(f)]
    
    _cached_index = (vectors, chunks)
    return _cached_index


def retrieve(query: str, top_k: int = TOP_K) -> List[dict]:
    """
    Runtime retrieval over the pre-built index.
    Returns up to top_k matching chunks as
    {crop, section, text, source, score} sorted by descending similarity.
    Chunks below RETRIEVAL_MIN_SIMILARITY are dropped — that's the
    "not in the dataset" signal.
    """
    vectors, chunks = load_index()
    q_vec = _embed([query])[0]
    sims = vectors @ q_vec  # normalised vectors: dot product == cosine similarity
    ranked = np.argsort(-sims)[:top_k]
    results = [chunks[int(i)] for i in ranked if sims[int(i)] >= RETRIEVAL_MIN_SIMILARITY]
    return [
        {
            "crop": c.crop,
            "section": c.section,
            "text": c.text[:MAX_CHUNK_CHARS],
            "source": c.source,
            "score": round(float(sims[ranked[idx]]), 3),
        }
        for idx, c in enumerate(results)
    ]


if __name__ == "__main__":
    cmd = sys.argv[1] if len(sys.argv) > 1 else "build"
    if cmd == "build":
        build_index()
    else:
        question = " ".join(sys.argv[1:])
        for m in retrieve(question):
            print(f"[{m['crop']} — {m['section']}] score={m['score']}")
            print(f"  {m['text'][:200]}...\n")
